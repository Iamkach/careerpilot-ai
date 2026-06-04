#!/usr/bin/env python3
"""
make_resume_template.py — generate a starter resume_template.docx
──────────────────────────────────────────────────────────────────
Builds a clean, single-column Word template with the docxtpl/Jinja2
placeholders that stage 2 fills in. This is a REFERENCE/starter — open
the result in Word and restyle fonts, colours, and spacing however you
like. As long as you keep the {{ ... }} / {% ... %} tags, tailoring will
keep working and every tailored resume will share this exact layout.

Run:  python scripts/make_resume_template.py
      python scripts/make_resume_template.py --out config/resume_template.docx
"""

import argparse
from pathlib import Path


def build_template(out_path: str) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Tighten global spacing so lines sit close together ──────
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(2)
    npf.line_spacing = 1.0

    # Compact section headings (small space above, none below)
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(12)
    hpf = h1.paragraph_format
    hpf.space_before = Pt(6)
    hpf.space_after = Pt(1)
    hpf.line_spacing = 1.0

    # Compact bullets
    try:
        bpf = doc.styles["List Bullet"].paragraph_format
        bpf.space_before = Pt(0)
        bpf.space_after = Pt(1)
        bpf.line_spacing = 1.0
    except KeyError:
        pass

    def tight(p, space_after=2):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)
        pf.line_spacing = 1.0
        return p

    # Header — name + contact line
    name = tight(doc.add_paragraph(), space_after=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("{{ name }}")
    run.bold = True
    run.font.size = Pt(20)

    contact = tight(doc.add_paragraph(), space_after=2)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("{{ contact }}").font.size = Pt(10)

    # Summary
    doc.add_heading("SUMMARY", level=1)
    tight(doc.add_paragraph("{{ summary }}"))

    # Skills — inline, bullet-separated
    doc.add_heading("SKILLS", level=1)
    tight(doc.add_paragraph(
        "{% for s in skills %}{{ s }}{% if not loop.last %} • {% endif %}{% endfor %}"
    ))

    # Experience — repeat block per role.
    # {%p ... %} = paragraph-level tag: docxtpl deletes the whole paragraph
    # holding the tag, so loop control adds NO blank lines to the output.
    doc.add_heading("EXPERIENCE", level=1)
    tight(doc.add_paragraph("{%p for job in experience %}"))
    role = tight(doc.add_paragraph())
    role.add_run("{{ job.company }} | {{ job.title }} | {{ job.dates }}").bold = True
    tight(doc.add_paragraph("{%p for b in job.bullets %}"))
    tight(doc.add_paragraph("{{ b }}", style="List Bullet"), space_after=1)
    tight(doc.add_paragraph("{%p endfor %}"))
    tight(doc.add_paragraph("{%p endfor %}"))

    # Education — repeat block
    doc.add_heading("EDUCATION", level=1)
    tight(doc.add_paragraph("{%p for ed in education %}"))
    tight(doc.add_paragraph("{{ ed.institution }} | {{ ed.degree }} | {{ ed.year }}"))
    tight(doc.add_paragraph("{%p endfor %}"))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="config/resume_template.docx",
                        help="Where to write the template .docx")
    args = parser.parse_args()
    path = build_template(args.out)
    print(f"[OK] Wrote starter template: {path}")
    print("  Open it in Word to restyle. Keep the {{ }} / {% %} tags intact.")
