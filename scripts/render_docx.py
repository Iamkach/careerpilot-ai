#!/usr/bin/env python3
"""
render_docx.py — in-place .docx text edits for tailored resumes
────────────────────────────────────────────────────────────────
Primary role today: `extract_docx_text()` / `apply_docx_edits()` apply targeted
`{old → new}` text replacements directly to a copy of the base resume `.docx`
(`RESUME_TEMPLATE_PATH`, default `config/Achyuth_Resume.docx`) via python-docx,
preserving the original document's formatting (fonts, spacing, run-level
bold/italic, ...) — this is the default flow stage 2 uses.

`render_resume_docx()` below is a legacy Jinja2/docxtpl template-render path
(paired with `scripts/make_resume_template.py`'s scaffolded
`config/resume_template.docx`) kept for reference but no longer used by the
default flow; it requires the optional `docxtpl` package to be installed
separately.

Expected `data` schema for `render_resume_docx()` (see RESUME_SCHEMA_DOC below):
    {
      "name": str,
      "contact": str,                 # single contact line
      "summary": str,
      "skills": [str, ...],
      "experience": [
        {"company": str, "title": str, "dates": str, "bullets": [str, ...]}
      ],
      "education": [
        {"institution": str, "degree": str, "year": str}
      ]
    }

Template placeholders (Jinja2 syntax inside the .docx):
    {{ name }}
    {{ contact }}
    {{ summary }}
    Skills line: {% for s in skills %}{{ s }}{% if not loop.last %} • {% endif %}{% endfor %}
    Experience: {% for job in experience %} ... {{ job.company }} | {{ job.title }} | {{ job.dates }}
                {% for b in job.bullets %}{{ b }}{% endfor %} ... {% endfor %}
    Education:  {% for ed in education %}{{ ed.institution }} | {{ ed.degree }} | {{ ed.year }}{% endfor %}
"""

import shutil
from pathlib import Path

from config.settings import YOUR_NAME

# Human-readable reference of the placeholders a template must use.
RESUME_SCHEMA_DOC = """\
Your config/resume_template.docx should contain these docxtpl (Jinja2) tags:

  {{ name }}                 — full name (header)
  {{ contact }}              — single contact line (email | linkedin | location)
  {{ summary }}              — professional summary paragraph

  Skills (inline, bullet-separated):
    {% for s in skills %}{{ s }}{% if not loop.last %} • {% endif %}{% endfor %}

  Experience (repeat block per role):
    {% for job in experience %}
    {{ job.company }} | {{ job.title }} | {{ job.dates }}
    {% for b in job.bullets %}
    {{ b }}
    {% endfor %}
    {% endfor %}

  Education (repeat block):
    {% for ed in education %}
    {{ ed.institution }} | {{ ed.degree }} | {{ ed.year }}
    {% endfor %}
"""

_REQUIRED_KEYS = ("name", "contact", "summary", "skills", "experience", "education")


def normalize_resume_data(data: dict) -> dict:
    """Fill in missing keys with safe empties so the template never crashes."""
    out = {
        "name": str(data.get("name", "") or ""),
        "contact": str(data.get("contact", "") or ""),
        "summary": str(data.get("summary", "") or ""),
        "skills": list(data.get("skills", []) or []),
        "experience": [],
        "education": [],
    }
    for job in data.get("experience", []) or []:
        out["experience"].append({
            "company": str(job.get("company", "") or ""),
            "title":   str(job.get("title", "") or ""),
            "dates":   str(job.get("dates", "") or ""),
            "bullets": list(job.get("bullets", []) or []),
        })
    for ed in data.get("education", []) or []:
        out["education"].append({
            "institution": str(ed.get("institution", "") or ""),
            "degree":      str(ed.get("degree", "") or ""),
            "year":        str(ed.get("year", "") or ""),
        })
    return out


def render_resume_docx(data: dict, template_path: str, out_path: str) -> str:
    """Render `data` into `template_path` and write the result to `out_path`.

    Raises FileNotFoundError if the template is missing, with guidance on
    the placeholders the template must contain.
    """
    try:
        from docxtpl import DocxTemplate
    except ImportError as e:
        raise ImportError(
            "docxtpl is required to render .docx resumes. "
            "Install it with: pip install docxtpl"
        ) from e

    tpl_path = Path(template_path)
    if not tpl_path.exists():
        raise FileNotFoundError(
            f"Resume template not found at {template_path}.\n\n{RESUME_SCHEMA_DOC}"
        )

    context = normalize_resume_data(data)

    doc = DocxTemplate(str(tpl_path))
    doc.render(context)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def extract_docx_text(docx_path: str) -> str:
    """Return all non-empty paragraph text from a .docx, one line per paragraph."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required. Install with: pip install python-docx") from e
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _apply_para_edits(para, replacements: list) -> None:
    """Apply one or more non-overlapping `(start, end, new)` replacements — offsets into
    the paragraph's ORIGINAL text — to a paragraph in a single pass, touching only the
    run(s) each span overlaps.

    Paragraph-level formatting (indent, bullet style, spacing) is preserved, as is each
    run's own formatting (bold, italic, ...) — a run entirely outside every matched span is
    never modified. Each replacement's new text is written into the first run its span
    overlaps, so it inherits that run's style rather than the paragraph's first run.

    Resolving every span against one original-text/run-layout snapshot up front (rather
    than re-scanning the paragraph's live text after each replacement) is what lets
    multiple edits share a paragraph safely: a later edit's span can never be thrown off by
    text an earlier edit already spliced in, and a run touched by two different spans gets
    both applied correctly instead of the second clobbering the first. `replacements` must
    already be non-overlapping (callers are responsible for resolving conflicts).
    """
    runs = para.runs
    if not runs:
        full = para.text
        pieces, pos = [], 0
        for start, end, new in replacements:
            pieces.append(full[pos:start])
            pieces.append(new)
            pos = end
        pieces.append(full[pos:])
        para.add_run("".join(pieces))
        return

    orig_texts = [r.text for r in runs]
    bounds = []
    pos = 0
    for text in orig_texts:
        bounds.append((pos, pos + len(text)))
        pos += len(text)

    # Bucket each replacement span into the run(s) (in original layout) it overlaps,
    # converting to run-local offsets. Only the first overlapping run for a given span
    # receives its `new` text; any further runs the same span crosses just drop their
    # overlapping slice (mirrors the single-edit "replaced" flag, generalized to N edits).
    per_run_patches = [[] for _ in runs]
    for start, end, new in replacements:
        first_run_in_span = True
        for i, (run_start, run_end) in enumerate(bounds):
            if run_end <= start or run_start >= end:
                continue  # entirely outside this span — leave untouched
            local_start = max(0, start - run_start)
            local_end = min(len(orig_texts[i]), end - run_start)
            per_run_patches[i].append((local_start, local_end, new if first_run_in_span else None))
            first_run_in_span = False

    for i, run in enumerate(runs):
        patches = per_run_patches[i]
        if not patches:
            continue
        patches.sort(key=lambda p: p[0])
        orig = orig_texts[i]
        pieces, pos = [], 0
        for local_start, local_end, new in patches:
            pieces.append(orig[pos:local_start])
            if new is not None:
                pieces.append(new)
            pos = local_end
        pieces.append(orig[pos:])
        run.text = "".join(pieces)


def _replace_para_text(para, old: str, new: str) -> None:
    """Replace the first occurrence of `old` with `new` inside a paragraph. Thin
    single-edit wrapper around `_apply_para_edits` — kept for callers/tests that only ever
    need one span resolved against the paragraph's current text."""
    start = para.text.find(old)
    if start == -1:
        return  # caller is expected to have already checked `old in para.text`
    _apply_para_edits(para, [(start, start + len(old), new)])


def apply_docx_edits(base_path: str, edits: list, out_path: str, job: dict | None = None) -> tuple[str, list]:
    """Copy base_path to out_path and apply targeted text replacements.

    edits: list of {"old": <exact existing text>, "new": <replacement text>}
    Each edit matches the first paragraph whose text contains the "old" string.

    job: optional {"title": str, "company": str} used to set per-resume docx
    core metadata (title/subject/keywords) below. When omitted, metadata still
    gets author/last_modified_by but falls back to a generic title/subject.

    Returns (out_path, unmatched_edits) — unmatched_edits lists any edit whose
    "old" text wasn't found in any paragraph, so the caller can surface it
    instead of it silently vanishing.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required. Install with: pip install python-docx") from e

    if not Path(base_path).exists():
        raise FileNotFoundError(
            f"Base resume not found at {base_path}. "
            "Set RESUME_TEMPLATE_PATH to an existing .docx file in config/settings.py."
        )

    shutil.copy2(base_path, out_path)
    doc = Document(out_path)

    paragraphs = doc.paragraphs
    # Snapshot every paragraph's ORIGINAL text once, up front, and resolve every edit's
    # match position against that fixed snapshot rather than the document's live (possibly
    # already-edited-by-an-earlier-edit) text. This is what prevents two edits landing in
    # the same paragraph from corrupting each other: if edit A's `new` text happens to
    # reintroduce a substring edit B's `old` matches, B still resolves against the
    # pre-edit text and can never mistake A's insertion for its own target.
    original_texts = [p.text for p in paragraphs]

    # para_index -> list of (start, end, new, edit) accepted so far for that paragraph.
    per_para_replacements: dict = {}
    unmatched = []
    for edit in edits:
        old = (edit.get("old") or "").strip()
        new = (edit.get("new") or "").strip()
        if not old or not new or old == new:
            continue

        target_idx = None
        start = -1
        for idx, text in enumerate(original_texts):
            pos = text.find(old)
            if pos != -1:
                target_idx = idx
                start = pos
                break

        if target_idx is None:
            unmatched.append(edit)
            continue

        end = start + len(old)
        bucket = per_para_replacements.setdefault(target_idx, [])
        # A genuine overlap between two edits' spans in the ORIGINAL text has no correct
        # resolution — skip the later one (report it as unmatched) rather than guess.
        if any(start < e and end > s for (s, e, _n, _ed) in bucket):
            unmatched.append(edit)
            continue
        bucket.append((start, end, new, edit))

    for idx, replacements in per_para_replacements.items():
        replacements.sort(key=lambda r: r[0])
        _apply_para_edits(paragraphs[idx], [(s, e, n) for (s, e, n, _ed) in replacements])

    job_title = (job or {}).get("title") or ""
    job_company = (job or {}).get("company") or ""
    props = doc.core_properties
    props.author = YOUR_NAME
    props.last_modified_by = YOUR_NAME
    props.category = "Resume"
    props.comments = ""
    props.title = f"{YOUR_NAME} - Resume" + (f" - {job_title}" if job_title else "")
    props.subject = job_title
    props.keywords = ", ".join(filter(None, [job_title, job_company]))

    doc.save(out_path)
    return out_path, unmatched


def resume_data_to_text(data: dict) -> str:
    """Flatten structured resume data into plain text (for quick review / .txt)."""
    d = normalize_resume_data(data)
    lines = [d["name"], d["contact"], ""]
    if d["summary"]:
        lines += ["SUMMARY", d["summary"], ""]
    if d["skills"]:
        lines += ["SKILLS", " • ".join(d["skills"]), ""]
    if d["experience"]:
        lines.append("EXPERIENCE")
        for job in d["experience"]:
            header = " | ".join(p for p in (job["company"], job["title"], job["dates"]) if p)
            lines.append(header)
            lines += [f"- {b}" for b in job["bullets"]]
            lines.append("")
    if d["education"]:
        lines.append("EDUCATION")
        for ed in d["education"]:
            lines.append(" | ".join(p for p in (ed["institution"], ed["degree"], ed["year"]) if p))
        lines.append("")
    return "\n".join(lines).strip() + "\n"
