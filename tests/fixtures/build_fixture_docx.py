"""
tests/fixtures/build_fixture_docx.py — builds the Phase 2 golden-file .docx fixture in code.

Generated (not committed as an opaque binary) so its exact paragraph/run structure is
reviewable in a diff and the golden expectations in test_render_docx.py are traceable back to
exactly what produced them.
"""
from pathlib import Path

from docx import Document


def build_fixture_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Krishna Achyuth")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("")  # blank paragraph — extract_docx_text() must skip this
    doc.add_paragraph("Summary: Experienced backend engineer with strong AWS skills.")

    # Multi-run paragraph, one run bold — exercises the run-collapsing characterization test.
    runs_para = doc.add_paragraph()
    runs_para.add_run("Led ")
    bold_run = runs_para.add_run("critical")
    bold_run.bold = True
    runs_para.add_run(" system migrations.")

    # "Python" appears once, "Java" appears once — sets up the same-paragraph double-edit
    # clobber characterization test (see test_render_docx.py).
    doc.add_paragraph("Skilled in Python and Java.")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
