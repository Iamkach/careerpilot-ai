"""
Phase 2 — golden-file tests for scripts/render_docx.py's extract_docx_text / apply_docx_edits.

No AI call happens in this module — the AI only produces the {old, new} edit list upstream in
stage 2 — so its I/O is fully deterministic given the fixture .docx, making it a strong
golden-file/snapshot-testing target. The fixture itself is built in code by
tests/fixtures/build_fixture_docx.py (see `fixture_docx_path` in conftest.py) rather than
committed as an opaque binary, so its exact structure stays reviewable in a diff.
"""
from pathlib import Path

import scripts.render_docx as render_docx
from scripts.render_docx import extract_docx_text, apply_docx_edits, convert_docx_to_pdf


# ── module docstring / python-docx dependency (PR #11 review item #7/#14) ──

def test_module_docstring_describes_python_docx_as_primary_path():
    # Importing the module at all proves python-docx is importable in this environment
    # (extract_docx_text/apply_docx_edits both `from docx import Document` internally) —
    # requirements.txt declaring python-docx is what makes that true on a fresh install.
    doc = render_docx.__doc__
    assert doc is not None
    assert "python-docx" in doc
    # The docstring must no longer present docxtpl as the module's primary/default role —
    # it's fine for the legacy render_resume_docx() path to still be mentioned by name.
    assert "render a tailored resume into a .docx template" not in doc
    first_line = doc.strip().splitlines()[0]
    assert "docxtpl" not in first_line.lower()


# ── extract_docx_text ────────────────────────────────────────────────────

def test_extract_docx_text_exact_expected_lines(fixture_docx_path):
    text = extract_docx_text(str(fixture_docx_path))
    assert text == (
        "Krishna Achyuth\n"
        "Senior Software Engineer\n"
        "Summary: Experienced backend engineer with strong AWS skills.\n"
        "Led critical system migrations.\n"
        "Skilled in Python and Java."
    )


def test_extract_docx_text_skips_blank_paragraphs(fixture_docx_path):
    # The fixture has a blank paragraph between the title lines and the summary — it must
    # not produce an empty line or extra blank entry in the joined output.
    text = extract_docx_text(str(fixture_docx_path))
    assert "\n\n" not in text


# ── apply_docx_edits — happy path ───────────────────────────────────────

def test_apply_docx_edits_applies_clean_non_overlapping_edits(fixture_docx_path, tmp_path):
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Senior Software Engineer", "new": "Staff Software Engineer"},
        {"old": "AWS skills", "new": "AWS and GCP skills"},
    ]
    result_path, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert result_path == str(out_path)
    assert unmatched == []
    text = extract_docx_text(str(out_path))
    assert "Staff Software Engineer" in text
    assert "Senior Software Engineer" not in text
    assert "AWS and GCP skills" in text


def test_apply_docx_edits_only_touches_the_matched_paragraph(fixture_docx_path, tmp_path):
    # The real requirement behind apply_docx_edits: a targeted ATS keyword edit must land
    # only in the one paragraph it's meant for and leave every other paragraph byte-for-byte
    # unchanged — a tailoring run touching unrelated resume content would be a real bug.
    out_path = tmp_path / "edited.docx"
    before = extract_docx_text(str(fixture_docx_path)).splitlines()

    apply_docx_edits(str(fixture_docx_path),
                      [{"old": "Senior Software Engineer", "new": "Staff Software Engineer"}],
                      str(out_path))

    after = extract_docx_text(str(out_path)).splitlines()
    assert len(after) == len(before)
    changed = [i for i, (b, a) in enumerate(zip(before, after)) if b != a]
    assert changed == [1]  # only the "Senior Software Engineer" line (index 1) changed
    assert after[1] == "Staff Software Engineer"


def test_apply_docx_edits_ignores_noop_and_blank_edits(fixture_docx_path, tmp_path):
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Krishna Achyuth", "new": "Krishna Achyuth"},  # no-op: old == new
        {"old": "", "new": "should be ignored"},               # blank old
        {"old": "should be ignored too", "new": ""},            # blank new
    ]
    _, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))
    # None of these are even attempted, so none show up as unmatched either.
    assert unmatched == []
    assert extract_docx_text(str(out_path)) == extract_docx_text(str(fixture_docx_path))


def test_apply_docx_edits_unmatched_edit_is_reported_not_silently_dropped(fixture_docx_path, tmp_path):
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "text that does not appear anywhere", "new": "replacement"},
    ]
    _, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert len(unmatched) == 1
    assert unmatched[0]["old"] == "text that does not appear anywhere"


def test_apply_docx_edits_writes_core_properties(fixture_docx_path, tmp_path):
    from docx import Document
    out_path = tmp_path / "edited.docx"
    job = {"title": "Senior Backend Engineer", "company": "Acme Corp"}

    apply_docx_edits(str(fixture_docx_path), [], str(out_path), job=job)

    props = Document(str(out_path)).core_properties
    assert "Senior Backend Engineer" in props.title
    assert props.subject == "Senior Backend Engineer"
    assert props.keywords == "Senior Backend Engineer, Acme Corp"


def test_apply_docx_edits_core_properties_fallback_without_job(fixture_docx_path, tmp_path):
    from docx import Document
    out_path = tmp_path / "edited.docx"

    apply_docx_edits(str(fixture_docx_path), [], str(out_path), job=None)

    props = Document(str(out_path)).core_properties
    assert props.subject == ""
    assert props.keywords == ""


# ── apply_docx_edits — run-level formatting preservation ────────────────

def test_apply_docx_edits_preserves_bold_run_on_mid_run_edit(fixture_docx_path, tmp_path):
    """A mid-run edit must not flatten the paragraph's other runs into the matched run's
    style. Fixture paragraph is "Led " (normal) + "critical" (bold) + " system migrations."
    (normal) — replacing "critical" must keep the bold run bold and leave the surrounding
    normal runs untouched and un-bold."""
    from docx import Document
    out_path = tmp_path / "edited.docx"
    edits = [{"old": "critical", "new": "highly critical and complex"}]

    apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    doc = Document(str(out_path))
    edited_para = next(p for p in doc.paragraphs if "highly critical and complex" in p.text)
    assert edited_para.text == "Led highly critical and complex system migrations."

    bold_runs = [r for r in edited_para.runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "highly critical and complex"

    non_bold_text = "".join(r.text for r in edited_para.runs if not r.bold)
    assert non_bold_text == "Led  system migrations."


def test_apply_docx_edits_preserves_formatting_when_edit_spans_run_boundary(fixture_docx_path, tmp_path):
    """An edit whose "old" text crosses a run boundary (here: the bold "critical" run and
    part of the following normal run) must still produce correct text, with the runs before
    the match left fully untouched."""
    from docx import Document
    out_path = tmp_path / "edited.docx"
    edits = [{"old": "critical system", "new": "urgent infrastructure"}]

    apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    doc = Document(str(out_path))
    edited_para = next(p for p in doc.paragraphs if "urgent infrastructure" in p.text)
    assert edited_para.text == "Led urgent infrastructure migrations."
    # The leading "Led " run is entirely before the match — must survive untouched.
    assert edited_para.runs[0].text == "Led "
    assert not edited_para.runs[0].bold


def test_apply_docx_edits_sequential_same_paragraph_edits_dont_clobber(fixture_docx_path, tmp_path):
    """Two edits landing in the same paragraph must each resolve against the paragraph's
    ORIGINAL text, not get applied sequentially against each other's output. Previously,
    edit A's `new` text reintroducing edit B's `old` substring caused B to clobber A's
    freshly-inserted text instead of the original occurrence it was meant for.

    Fixture paragraph: "Skilled in Python and Java."
      Edit A: old="Python", new="Java"   → intends "Skilled in Java and Java."
      Edit B: old="Java",   new="Kotlin" → must target the ORIGINAL "Java" at the end, not
              the one edit A introduces.
    Correct result: "Skilled in Java and Kotlin." — both edits land on their own,
    independently-resolved occurrence."""
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Python", "new": "Java"},
        {"old": "Java", "new": "Kotlin"},
    ]

    result_path, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert unmatched == []
    text = extract_docx_text(str(out_path))
    assert "Skilled in Java and Kotlin." in text
    assert "Skilled in Kotlin and Java." not in text  # the old, buggy result


def test_apply_docx_edits_same_paragraph_edit_order_reversed_still_correct(fixture_docx_path, tmp_path):
    """Same scenario as above but with the edits given in the opposite order, to confirm the
    fix isn't relying on edit A happening to be processed (and thus resolved) first."""
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Java", "new": "Kotlin"},
        {"old": "Python", "new": "Java"},
    ]

    _, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert unmatched == []
    text = extract_docx_text(str(out_path))
    assert "Skilled in Java and Kotlin." in text


def test_apply_docx_edits_same_paragraph_earlier_edit_shifts_length_before_later_one(fixture_docx_path, tmp_path):
    """An earlier edit in the same paragraph that changes text length (here, a much longer
    replacement) must not throw off where the later edit's match is applied — positions are
    resolved against the original text up front, so length deltas from earlier edits in the
    same paragraph never shift a later edit's target."""
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Python", "new": "a very long and verbose Python-adjacent phrase"},
        {"old": "Java", "new": "Kotlin"},
    ]

    _, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert unmatched == []
    text = extract_docx_text(str(out_path))
    assert "Skilled in a very long and verbose Python-adjacent phrase and Kotlin." in text


def test_apply_docx_edits_genuinely_overlapping_same_paragraph_edits_second_is_skipped(fixture_docx_path, tmp_path):
    """Two edits whose `old` substrings genuinely overlap in the ORIGINAL paragraph text have
    no correct simultaneous resolution — the second (in edit-list order) must be reported as
    unmatched rather than corrupting the paragraph, and the first must still apply cleanly."""
    out_path = tmp_path / "edited.docx"
    edits = [
        {"old": "Skilled in Python", "new": "Proficient in Python"},
        {"old": "in Python and", "new": "in Rust and"},  # overlaps the first edit's span
    ]

    _, unmatched = apply_docx_edits(str(fixture_docx_path), edits, str(out_path))

    assert len(unmatched) == 1
    assert unmatched[0]["old"] == "in Python and"
    text = extract_docx_text(str(out_path))
    assert "Proficient in Python and Java." in text


# ── table support (PR #11 review item #15) ─────────────────────────────

def _build_docx_with_table(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Krishna Achyuth")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("Python")
    table.cell(0, 1).paragraphs[0].add_run("Experienced with AWS and Docker.")
    path = tmp_path / "table_resume.docx"
    doc.save(str(path))
    return str(path)


def test_extract_docx_text_includes_table_cell_paragraphs(tmp_path):
    path = _build_docx_with_table(tmp_path)
    text = extract_docx_text(path)
    assert "Krishna Achyuth" in text
    assert "Python" in text
    assert "Experienced with AWS and Docker." in text


def test_apply_docx_edits_matches_text_inside_a_table_cell(tmp_path):
    path = _build_docx_with_table(tmp_path)
    out_path = tmp_path / "table_resume_edited.docx"
    edits = [{"old": "AWS and Docker", "new": "AWS, GCP, and Docker"}]

    result_path, unmatched = apply_docx_edits(path, edits, str(out_path))

    assert unmatched == []
    text = extract_docx_text(result_path)
    assert "Experienced with AWS, GCP, and Docker." in text


# ── convert_docx_to_pdf ──────────────────────────────────────────────────
# Step 10 residual gap #2: Stage 7's fill path (autoapply_browser._resolve_upload_path)
# needs a .docx→.pdf fallback for ATS forms that reject .docx. LibreOffice (`soffice`) is a
# system install, not a pip package, so every path here must degrade to None rather than
# raise when it's missing — mirroring sources._headless_fetch()'s optional-dependency
# contract — and none of these tests require LibreOffice to actually be installed.

def test_convert_docx_to_pdf_returns_none_when_soffice_missing(monkeypatch, tmp_path):
    monkeypatch.setattr(render_docx.shutil, "which", lambda name: None)
    docx = tmp_path / "resume.docx"
    docx.write_text("not a real docx, just needs to exist")
    assert convert_docx_to_pdf(str(docx)) is None


def test_convert_docx_to_pdf_returns_none_for_missing_source_file(monkeypatch, tmp_path):
    monkeypatch.setattr(render_docx.shutil, "which", lambda name: "/usr/bin/soffice")
    assert convert_docx_to_pdf(str(tmp_path / "does_not_exist.docx")) is None


def test_convert_docx_to_pdf_returns_none_when_subprocess_fails(monkeypatch, tmp_path):
    import subprocess as _subprocess

    monkeypatch.setattr(render_docx.shutil, "which", lambda name: "/usr/bin/soffice")

    def _raise(*args, **kwargs):
        raise _subprocess.CalledProcessError(1, "soffice")

    monkeypatch.setattr(render_docx.subprocess, "run", _raise)
    docx = tmp_path / "resume.docx"
    docx.write_text("not a real docx, just needs to exist")
    assert convert_docx_to_pdf(str(docx)) is None


def test_convert_docx_to_pdf_returns_pdf_path_on_success(monkeypatch, tmp_path):
    monkeypatch.setattr(render_docx.shutil, "which", lambda name: "/usr/bin/soffice")

    def _fake_run(cmd, **kwargs):
        # Simulate soffice actually writing the .pdf next to the .docx.
        out_dir = Path(cmd[cmd.index("--outdir") + 1])
        src = Path(cmd[-1])
        (out_dir / (src.stem + ".pdf")).write_bytes(b"%PDF-1.4 fake")
        return None

    monkeypatch.setattr(render_docx.subprocess, "run", _fake_run)
    docx = tmp_path / "resume.docx"
    docx.write_text("not a real docx, just needs to exist")

    result = convert_docx_to_pdf(str(docx))

    assert result == str(tmp_path / "resume.pdf")
    assert Path(result).exists()
